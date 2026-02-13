"""
Unit tests for Phase 1B: Broader File Format Support.

Tests BaseFileProcessor, FileProcessorFactory, ImageProcessor,
SpreadsheetProcessor, DocxProcessor, EDIProcessor, and DicomProcessor.
"""

import csv
import io
import tempfile
from pathlib import Path
from typing import Any

import pytest
from PIL import Image

from src.preprocessing.base_processor import (
    BaseFileProcessor,
    FileFormat,
    FileValidationError,
    IMAGE_FORMATS,
    SPREADSHEET_FORMATS,
    SUPPORTED_EXTENSIONS,
    UnsupportedFormatError,
)
from src.preprocessing.file_factory import FileProcessorFactory
from src.preprocessing.image_processor import ImageProcessor


# ──────────────────────────────────────────────────────────────────
# Helper fixtures
# ──────────────────────────────────────────────────────────────────


@pytest.fixture
def tmp_dir():
    """Provide a temporary directory."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def sample_png(tmp_dir: Path) -> Path:
    """Create a sample PNG file."""
    path = tmp_dir / "test.png"
    img = Image.new("RGB", (200, 100), color="white")
    img.save(str(path), format="PNG")
    return path


@pytest.fixture
def sample_jpg(tmp_dir: Path) -> Path:
    """Create a sample JPG file."""
    path = tmp_dir / "test.jpg"
    img = Image.new("RGB", (200, 100), color="blue")
    img.save(str(path), format="JPEG")
    return path


@pytest.fixture
def sample_tiff(tmp_dir: Path) -> Path:
    """Create a sample TIFF file."""
    path = tmp_dir / "test.tiff"
    img = Image.new("RGB", (200, 100), color="red")
    img.save(str(path), format="TIFF")
    return path


@pytest.fixture
def sample_bmp(tmp_dir: Path) -> Path:
    """Create a sample BMP file."""
    path = tmp_dir / "test.bmp"
    img = Image.new("RGB", (200, 100), color="green")
    img.save(str(path), format="BMP")
    return path


@pytest.fixture
def sample_csv(tmp_dir: Path) -> Path:
    """Create a sample CSV file."""
    path = tmp_dir / "test.csv"
    with path.open("w", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["Name", "Age", "City"])
        writer.writerow(["Alice", "30", "New York"])
        writer.writerow(["Bob", "25", "San Francisco"])
        writer.writerow(["Charlie", "35", "Chicago"])
    return path


@pytest.fixture
def sample_edi(tmp_dir: Path) -> Path:
    """Create a minimal EDI 837 file."""
    path = tmp_dir / "test.837"
    # ISA segment is fixed 106 chars, element sep at pos 3, segment sep at pos 105
    content = (
        "ISA*00*          *00*          *ZZ*SENDER         *ZZ*RECEIVER       *"
        "240101*1200*^*00501*000000001*0*P*:~"
        "GS*HC*SENDER*RECEIVER*20240101*1200*1*X*005010X222A1~"
        "ST*837*0001~"
        "BHT*0019*00*123456*20240101*1200*CH~"
        "NM1*41*2*BILLING PROVIDER*****46*123456789~"
        "CLM*CLAIM001*500***11:B:1*Y*A*Y*Y~"
        "SE*6*0001~"
        "GE*1*1~"
        "IEA*1*000000001~"
    )
    path.write_text(content, encoding="utf-8")
    return path


# ──────────────────────────────────────────────────────────────────
# FileFormat Enum Tests
# ──────────────────────────────────────────────────────────────────


class TestFileFormat:
    """Tests for FileFormat enum."""

    def test_pdf_format(self):
        assert FileFormat.PDF.value == "pdf"

    def test_image_formats(self):
        assert FileFormat.PNG in IMAGE_FORMATS
        assert FileFormat.JPG in IMAGE_FORMATS
        assert FileFormat.JPEG in IMAGE_FORMATS
        assert FileFormat.TIFF in IMAGE_FORMATS
        assert FileFormat.BMP in IMAGE_FORMATS

    def test_spreadsheet_formats(self):
        assert FileFormat.XLSX in SPREADSHEET_FORMATS
        assert FileFormat.CSV in SPREADSHEET_FORMATS

    def test_supported_extensions_complete(self):
        """All formats have corresponding extensions."""
        for fmt in FileFormat:
            assert f".{fmt.value}" in SUPPORTED_EXTENSIONS


# ──────────────────────────────────────────────────────────────────
# FileProcessorFactory Tests
# ──────────────────────────────────────────────────────────────────


class TestFileProcessorFactory:
    """Tests for FileProcessorFactory routing."""

    def test_factory_routes_png(self, sample_png: Path):
        factory = FileProcessorFactory()
        processor = factory.get_processor(sample_png)
        assert isinstance(processor, ImageProcessor)

    def test_factory_routes_jpg(self, sample_jpg: Path):
        factory = FileProcessorFactory()
        processor = factory.get_processor(sample_jpg)
        assert isinstance(processor, ImageProcessor)

    def test_factory_routes_csv(self, sample_csv: Path):
        from src.preprocessing.spreadsheet_processor import SpreadsheetProcessor
        factory = FileProcessorFactory()
        processor = factory.get_processor(sample_csv)
        assert isinstance(processor, SpreadsheetProcessor)

    def test_factory_routes_edi(self, sample_edi: Path):
        from src.preprocessing.edi_processor import EDIProcessor
        factory = FileProcessorFactory()
        processor = factory.get_processor(sample_edi)
        assert isinstance(processor, EDIProcessor)

    def test_factory_rejects_unsupported(self, tmp_dir: Path):
        path = tmp_dir / "test.xyz"
        path.write_text("unsupported")
        factory = FileProcessorFactory()
        with pytest.raises(UnsupportedFormatError):
            factory.get_processor(path)

    def test_factory_is_supported(self, sample_png: Path, tmp_dir: Path):
        factory = FileProcessorFactory()
        assert factory.is_supported(sample_png)
        unsupported = tmp_dir / "test.xyz"
        assert not factory.is_supported(unsupported)

    def test_factory_supported_extensions(self):
        exts = FileProcessorFactory.supported_extensions()
        assert ".pdf" in exts
        assert ".png" in exts
        assert ".csv" in exts

    def test_factory_caches_processors(self, sample_png: Path, sample_jpg: Path):
        """Same processor type is reused for same format category."""
        factory = FileProcessorFactory()
        p1 = factory.get_processor(sample_png)
        p2 = factory.get_processor(sample_jpg)
        # Both image formats use the same cached ImageProcessor instance
        assert p1 is p2


# ──────────────────────────────────────────────────────────────────
# ImageProcessor Tests
# ──────────────────────────────────────────────────────────────────


class TestImageProcessor:
    """Tests for ImageProcessor."""

    def test_process_png(self, sample_png: Path):
        processor = ImageProcessor()
        result = processor.process(sample_png)
        assert result.page_count == 1
        assert result.pages[0].page_number == 1
        assert result.pages[0].width == 200
        assert result.pages[0].height == 100
        assert result.metadata.file_name == "test.png"

    def test_process_jpg(self, sample_jpg: Path):
        processor = ImageProcessor()
        result = processor.process(sample_jpg)
        assert result.page_count == 1
        assert result.pages[0].image_bytes  # has content

    def test_process_tiff(self, sample_tiff: Path):
        processor = ImageProcessor()
        result = processor.process(sample_tiff)
        assert result.page_count == 1

    def test_process_bmp(self, sample_bmp: Path):
        processor = ImageProcessor()
        result = processor.process(sample_bmp)
        assert result.page_count == 1

    def test_process_large_image_resized(self, tmp_dir: Path):
        """Large images are resized to max dimension."""
        path = tmp_dir / "large.png"
        img = Image.new("RGB", (5000, 3000), "white")
        img.save(str(path), format="PNG")

        processor = ImageProcessor(max_dimension=2048)
        result = processor.process(path)
        assert result.pages[0].width <= 2048
        assert result.pages[0].height <= 2048
        assert len(result.warnings) > 0  # Should warn about resize

    def test_validate_missing_file(self):
        processor = ImageProcessor()
        with pytest.raises(FileValidationError):
            processor.validate(Path("/nonexistent/file.png"))

    def test_page_image_has_base64(self, sample_png: Path):
        processor = ImageProcessor()
        result = processor.process(sample_png)
        page = result.pages[0]
        assert page.base64_encoded
        assert page.data_uri.startswith("data:image/png;base64,")

    def test_page_image_has_metadata(self, sample_png: Path):
        processor = ImageProcessor()
        result = processor.process(sample_png)
        assert result.metadata.file_hash
        assert result.metadata.processing_id.startswith("img_")
        assert result.processing_time_ms >= 0

    def test_process_via_factory(self, sample_png: Path):
        """End-to-end: factory → ImageProcessor → ProcessingResult."""
        factory = FileProcessorFactory()
        result = factory.process(sample_png)
        assert result.page_count == 1
        assert result.pages[0].width == 200


# ──────────────────────────────────────────────────────────────────
# SpreadsheetProcessor Tests
# ──────────────────────────────────────────────────────────────────


class TestSpreadsheetProcessor:
    """Tests for SpreadsheetProcessor."""

    def test_process_csv(self, sample_csv: Path):
        from src.preprocessing.spreadsheet_processor import SpreadsheetProcessor
        processor = SpreadsheetProcessor()
        result = processor.process(sample_csv)
        assert result.page_count >= 1
        # Text content should contain the data
        text = result.pages[0].text_content
        assert "Alice" in text or "Name" in text

    def test_csv_metadata(self, sample_csv: Path):
        from src.preprocessing.spreadsheet_processor import SpreadsheetProcessor
        processor = SpreadsheetProcessor()
        result = processor.process(sample_csv)
        assert result.metadata.file_name == "test.csv"
        assert result.metadata.processing_id.startswith("sheet_")

    def test_validate_missing_file(self):
        from src.preprocessing.spreadsheet_processor import SpreadsheetProcessor
        processor = SpreadsheetProcessor()
        with pytest.raises(FileValidationError):
            processor.validate(Path("/nonexistent/file.csv"))

    def test_process_csv_via_factory(self, sample_csv: Path):
        factory = FileProcessorFactory()
        result = factory.process(sample_csv)
        assert result.page_count >= 1


# ──────────────────────────────────────────────────────────────────
# EDIProcessor Tests
# ──────────────────────────────────────────────────────────────────


class TestEDIProcessor:
    """Tests for EDIProcessor."""

    def test_process_edi(self, sample_edi: Path):
        from src.preprocessing.edi_processor import EDIProcessor
        processor = EDIProcessor()
        result = processor.process(sample_edi)
        assert result.page_count >= 1
        # Should detect 837 transaction type
        assert "837" in (result.metadata.title or "")

    def test_edi_text_content(self, sample_edi: Path):
        from src.preprocessing.edi_processor import EDIProcessor
        processor = EDIProcessor()
        result = processor.process(sample_edi)
        text = result.pages[0].text_content
        assert "Claim Information" in text or "CLM" in text

    def test_edi_metadata(self, sample_edi: Path):
        from src.preprocessing.edi_processor import EDIProcessor
        processor = EDIProcessor()
        result = processor.process(sample_edi)
        assert result.metadata.processing_id.startswith("edi_")

    def test_validate_non_edi_file(self, tmp_dir: Path):
        from src.preprocessing.edi_processor import EDIProcessor
        path = tmp_dir / "bad.edi"
        path.write_text("This is not EDI content", encoding="utf-8")
        processor = EDIProcessor()
        with pytest.raises(FileValidationError, match="ISA"):
            processor.validate(path)

    def test_process_edi_via_factory(self, sample_edi: Path):
        factory = FileProcessorFactory()
        result = factory.process(sample_edi)
        assert result.page_count >= 1


# ──────────────────────────────────────────────────────────────────
# DocxProcessor Tests (requires python-docx)
# ──────────────────────────────────────────────────────────────────


class TestDocxProcessor:
    """Tests for DocxProcessor (skipped if python-docx not installed)."""

    @pytest.fixture
    def sample_docx(self, tmp_dir: Path) -> Path | None:
        """Create a sample DOCX if python-docx is available."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
            return None

        path = tmp_dir / "test.docx"
        doc = Document()
        doc.add_paragraph("Patient Name: John Doe")
        doc.add_paragraph("Date of Birth: 01/15/1990")
        doc.add_paragraph("Diagnosis: Type 2 Diabetes")
        doc.save(str(path))
        return path

    def test_process_docx(self, sample_docx: Path):
        if sample_docx is None:
            pytest.skip("python-docx not installed")
        from src.preprocessing.docx_processor import DocxProcessor
        processor = DocxProcessor()
        result = processor.process(sample_docx)
        assert result.page_count >= 1
        assert "John Doe" in result.pages[0].text_content

    def test_docx_metadata(self, sample_docx: Path):
        if sample_docx is None:
            pytest.skip("python-docx not installed")
        from src.preprocessing.docx_processor import DocxProcessor
        processor = DocxProcessor()
        result = processor.process(sample_docx)
        assert result.metadata.processing_id.startswith("docx_")

    def test_process_docx_via_factory(self, sample_docx: Path):
        if sample_docx is None:
            pytest.skip("python-docx not installed")
        factory = FileProcessorFactory()
        result = factory.process(sample_docx)
        assert result.page_count >= 1
